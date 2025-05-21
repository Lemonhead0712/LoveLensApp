import docx
from docx import Document
from docx.shared import Inches
import json
import sys
import io
import base64
from PIL import Image
import matplotlib.pyplot as plt
import numpy as np

def generate_word_document(analysis_data):
    # Create a new Document
    doc = Document()
    
    # Add a title
    doc.add_heading('💖 Love Lens: Relationship Insight', 0)
    
    # Add a note to reader
    doc.add_heading('Note to Reader', level=1)
    doc.add_paragraph(
        'This is a third-party relationship reflection based on real conversations. '
        'The goal? Clarity. All emotional tones are preserved as they were sent. '
        'We're not assigning blame—just holding up a mirror to the emotional patterns at play.'
    )
    
    # Add Communication Styles section
    doc.add_heading('💬 Communication Styles & Emotional Tone', level=1)
    doc.add_paragraph(analysis_data['communicationStyles'])
    
    # Add Recurring Patterns section
    doc.add_heading('🔁 Recurring Patterns Identified', level=1)
    doc.add_paragraph(analysis_data['recurringPatterns'])
    
    # Add Reflective Frameworks section
    doc.add_heading('🧠 Reflective Frameworks', level=1)
    doc.add_paragraph(analysis_data['reflectiveFrameworks'])
    
    # Add What's Getting in the Way section
    doc.add_heading('🚧 What's Getting in the Way', level=1)
    doc.add_paragraph(analysis_data['gettingInTheWay'])
    
    # Add Constructive Feedback section
    doc.add_heading('🌱 Constructive Feedback', level=1)
    doc.add_paragraph(analysis_data['constructiveFeedback'])
    
    # Add charts
    doc.add_heading('📊 Visual Insights', level=1)
    
    # In a real implementation, we would generate charts using matplotlib
    # and add them to the document
    doc.add_paragraph('Charts would be generated and inserted here.')
    
    # Add Outlook section
    doc.add_heading('🔮 Outlook', level=1)
    doc.add_paragraph(analysis_data['outlook'])
    
    # Add Optional Appendix section
    doc.add_heading('📎 Optional Appendix', level=1)
    doc.add_paragraph(analysis_data['optionalAppendix'])
    
    # Save the document
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

# This function would be called from a server action or API route
def main():
    # Read input JSON from stdin
    analysis_data = json.loads(sys.stdin.read())
    
    # Generate the Word document
    doc_buffer = generate_word_document(analysis_data)
    
    # Return the document as base64
    print(base64.b64encode(doc_buffer.getvalue()).decode('utf-8'))

if __name__ == "__main__":
    main()
